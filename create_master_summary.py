#!/usr/bin/env python3
"""
Create MASTER_PROJECT_SUMMARY.docx for NOD2 Crohn's Disease Drug Discovery ISEF Project.

This script generates a comprehensive Word document summarizing all project findings,
with corrected compound identity (CID_10120 Bufadienolide, NOT CID_10592 Dihydrocortisol).

Output:
    - MASTER_PROJECT_SUMMARY.docx
    - MASTER_PROJECT_SUMMARY.pdf (if docx2pdf available)

Author: Tanmay Vasudeva
Project: ISEF 2026 - Fort Worth Regional Science Fair
"""

import json
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx")
    exit(1)

# Paths
BASE_DIR = Path(__file__).parent
OUTPUT_DOCX = BASE_DIR / "MASTER_PROJECT_SUMMARY.docx"
OUTPUT_PDF = BASE_DIR / "MASTER_PROJECT_SUMMARY.pdf"

# ============================================================================
# PROJECT METADATA
# ============================================================================
PROJECT_INFO = {
    "title": "Precision Medicine, N=1: An FEP-Validated Deep Learning Pipeline to Rescue the NOD2 R702W Crohn's Variant",
    "student": "Tanmay Vasudeva",
    "school": "Texas Virtual Academy at Hallsville (Harrison County, Texas)",
    "grade": "Junior (Class of 2027)",
    "competition": "Fort Worth Regional Science Fair",
    "competition_date": "February 21, 2026",
    "forms_deadline": "January 23, 2026",
    "start_date": "May 2025",
    "data_end": "January 2026",
}

# ============================================================================
# FEP RESULTS (CORRECTED VALUES)
# ============================================================================
FEP_RESULTS = {
    "febuxostat": {
        "type": "FDA Drug (Gout)",
        "pubchem_cid": "134018",
        "formula": "C16H16N2O3S",
        "mw": "316.37",
        "dG_bind_wt": "-10.36 +/- 0.18",
        "dG_bind_mut": "-8.02 +/- 0.19",
        "ddG": "+2.34",
        "Kd_wt": "~25 nM",
        "Kd_mut": "~1.4 uM",
        "mutation_effect": "50x WEAKER in mutant",
        "significance": "8 sigma (p<0.001)",
    },
    "cid_10120": {
        "type": "Natural Product (Bufadienolide)",
        "pubchem_cid": "10120",
        "formula": "C24H34O5",
        "mw": "402.53",
        "dG_bind_wt": "-15.22 +/- 0.26",
        "dG_bind_mut": "-15.66 +/- 0.26",
        "ddG": "-0.44",
        "Kd_wt": "~5 pM",
        "Kd_mut": "~3 pM",
        "mutation_effect": "NO EFFECT (mutation-resistant)",
        "significance": "1.8 sigma (not significant)",
    }
}

# ============================================================================
# MM-GBSA vs FEP COMPARISON (using absolute binding ΔΔG values)
# ============================================================================
MMGBSA_FEP_COMPARISON = [
    {"compound": "Febuxostat", "mmgbsa_ddg": "-2.7", "fep_ddg": "+2.34", "direction": "OPPOSITE SIGN"},
    {"compound": "Bufadienolide", "mmgbsa_ddg": "-3.9", "fep_ddg": "-0.44", "direction": "9x OVERESTIMATE"},
]

# ============================================================================
# PROJECT PHASES
# ============================================================================
PROJECT_PHASES = [
    {"phase": "0", "name": "Foundation", "status": "Complete", "description": "Literature review, gap analysis", "output": "Research plan"},
    {"phase": "1", "name": "Structure Prep", "status": "Complete", "description": "NOD2 LRR domain preparation", "output": "human_NOD2_LRR.pdbqt"},
    {"phase": "2", "name": "Library Curation", "status": "Complete", "description": "7,414 natural + 2,152 FDA compounds", "output": "9,566 total"},
    {"phase": "3", "name": "Virtual Screening", "status": "Complete", "description": "GNINA GPU-accelerated docking", "output": "Top 500 hits"},
    {"phase": "4", "name": "NOD2-Scout ML", "status": "Complete", "description": "Random Forest ranking model", "output": "0.90 AUC-ROC"},
    {"phase": "5", "name": "ADMET Filtering", "status": "Complete", "description": "Toxicity, drug-likeness", "output": "8 Tier-1 candidates"},
    {"phase": "6", "name": "MD Simulations (WT)", "status": "Complete", "description": "20ns x 14 trajectories", "output": "Pocket occupancy data"},
    {"phase": "7", "name": "MM-GBSA", "status": "Complete", "description": "Binding energy estimation", "output": "DDG estimates"},
    {"phase": "8", "name": "Cross-Validation", "status": "Complete", "description": "Chai-1, DiffDock, FlowDock", "output": "Confirmed binding site"},
    {"phase": "9", "name": "Mutation Analysis", "status": "Complete", "description": "Distance calculations for all mutations", "output": "R702W = 79.4 A from pocket"},
    {"phase": "10", "name": "MD Simulations (Mutant)", "status": "Complete", "description": "R702W + G908R, 240ns total", "output": "Trajectory data"},
    {"phase": "11", "name": "Clinical Trial Design", "status": "Complete", "description": "Phase II protocol", "output": "88% power, 210 patients"},
    {"phase": "12", "name": "FEP Validation", "status": "Complete", "description": "80 complex windows (40 Feb + 40 Buf)", "output": "DDG with error bars"},
    {"phase": "13", "name": "Absolute dG", "status": "Complete", "description": "Solvent legs, thermodynamic cycle", "output": "Final binding affinities"},
]

# ============================================================================
# STATISTICS
# ============================================================================
STATISTICS = {
    "total_compounds_screened": "9,566",
    "library_breakdown": "7,414 natural products + 2,152 FDA drugs",
    "ml_model_auc_roc": "0.90",
    "md_simulations_wt": "14 trajectories, 280 ns",
    "md_simulations_mutant": "12 trajectories, 240 ns",
    "total_simulation_time": "520 ns",
    "fep_windows_febuxostat": "40 complex (20 WT + 20 MUT) + 20 solvent = 60",
    "fep_windows_bufadienolide": "40 complex (20 WT + 20 MUT) + 20 solvent = 60",
    "solvent_windows": "40 total (20 per compound)",
    "total_fep_windows": "120 (80 complex + 40 solvent)",
}

# ============================================================================
# COMPOUND IDENTITY CORRECTION
# ============================================================================
COMPOUND_CORRECTION = {
    "old": {
        "cid": "CID_10592",
        "name": "5beta-Dihydrocortisol",
        "type": "Cortisol metabolite",
        "formula": "C21H32O5",
        "mw": "364.48",
    },
    "new": {
        "cid": "CID_10120",
        "name": "3beta,5,14-Trihydroxy-5beta-bufa-20,22-dienolide",
        "common_name": "Bufadienolide",
        "type": "Cardiac glycoside (toad/plant compound)",
        "formula": "C24H34O5",
        "mw": "402.53",
    },
    "root_cause": "When selecting the top natural product for MD simulation, the SDF file for the #1 CNN affinity compound (CID_10120 = 6.756) was grabbed instead of the ADMET-ranked compound (CID_10592 = 6.741). The label 'CID_10592' was incorrectly applied to all downstream work.",
}


def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading(doc, text, level=1):
    """Add a heading with consistent formatting."""
    heading = doc.add_heading(text, level=level)
    return heading


def add_table(doc, headers, rows, header_color="2E86AB"):
    """Add a formatted table."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'

    # Header row
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(hdr_cells[i], header_color)

    # Data rows
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = str(cell_data)

    doc.add_paragraph()  # Space after table
    return table


def create_document():
    """Create the master project summary document."""
    doc = Document()

    # ========================================================================
    # TITLE PAGE
    # ========================================================================
    title = doc.add_heading(PROJECT_INFO["title"], 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("MASTER PROJECT SUMMARY").bold = True

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run(f"\nStudent: {PROJECT_INFO['student']}\n")
    info.add_run(f"School: {PROJECT_INFO['school']}\n")
    info.add_run(f"Grade: {PROJECT_INFO['grade']}\n")
    info.add_run(f"\nCompetition: {PROJECT_INFO['competition']}\n")
    info.add_run(f"Date: {PROJECT_INFO['competition_date']}\n")
    info.add_run(f"\nGenerated: {datetime.now().strftime('%Y-%m-%d %H:%M')}\n")

    doc.add_page_break()

    # ========================================================================
    # TABLE OF CONTENTS (placeholder)
    # ========================================================================
    doc.add_heading("Table of Contents", 1)
    toc = doc.add_paragraph()
    toc.add_run("1. Executive Summary\n")
    toc.add_run("2. Project Phases\n")
    toc.add_run("3. Final FEP Results\n")
    toc.add_run("4. MM-GBSA vs FEP Discrepancy\n")
    toc.add_run("5. File Paths\n")
    toc.add_run("6. Compound Identity Investigation\n")
    toc.add_run("7. Statistics Summary\n")
    toc.add_run("8. Methods Detail\n")
    toc.add_run("9. Novel Contributions\n")
    toc.add_run("10. Future Directions\n")

    doc.add_page_break()

    # ========================================================================
    # SECTION 1: EXECUTIVE SUMMARY
    # ========================================================================
    doc.add_heading("1. Executive Summary", 1)

    doc.add_heading("Problem", 2)
    doc.add_paragraph(
        "NOD2 mutations are the #1 genetic risk factor for Crohn's disease, "
        "increasing risk up to 40x. No drugs specifically target NOD2."
    )

    doc.add_heading("Approach", 2)
    doc.add_paragraph(
        "Built a computational drug discovery pipeline combining virtual screening, "
        "machine learning, molecular dynamics, and gold-standard free energy calculations."
    )

    doc.add_heading("Key Results", 2)
    results = doc.add_paragraph()
    results.add_run("- Screened 9,566 compounds\n")
    results.add_run("- Built NOD2-Scout ML model (AUC-ROC = 0.90)\n")
    results.add_run("- Validated 2 compounds with FEP:\n")
    results.add_run("  1. ").bold = True
    results.add_run("Febuxostat").bold = True
    results.add_run(" (FDA-approved gout drug): Binds NOD2 but 50x weaker in R702W mutant\n")
    results.add_run("  2. ").bold = True
    results.add_run("CID_10120 Bufadienolide").bold = True
    results.add_run(" (natural product): Binds orders of magnitude stronger than febuxostat (predicted) AND mutation-resistant in silico\n")

    doc.add_heading("Novel Discovery", 2)
    novel = doc.add_paragraph()
    novel.add_run("First ever report of bufadienolide-NOD2 binding. ").bold = True
    novel.add_run(
        "This class of compounds are known NF-kB inhibitors, and NOD2 signals through NF-kB - "
        "suggesting a dual mechanism."
    )

    doc.add_page_break()

    # ========================================================================
    # SECTION 2: PROJECT PHASES
    # ========================================================================
    doc.add_heading("2. Project Phases", 1)

    headers = ["Phase", "Name", "Status", "Description", "Key Output"]
    rows = [[p["phase"], p["name"], p["status"], p["description"], p["output"]] for p in PROJECT_PHASES]
    add_table(doc, headers, rows)

    doc.add_page_break()

    # ========================================================================
    # SECTION 3: FINAL FEP RESULTS
    # ========================================================================
    doc.add_heading("3. Final FEP Results", 1)

    doc.add_heading("FEP-Validated NOD2 Binders", 2)

    feb = FEP_RESULTS["febuxostat"]
    buf = FEP_RESULTS["cid_10120"]

    headers = ["Property", "Febuxostat", "CID_10120 (Bufadienolide)"]
    rows = [
        ["Type", feb["type"], buf["type"]],
        ["PubChem CID", feb["pubchem_cid"], buf["pubchem_cid"]],
        ["Formula", feb["formula"], buf["formula"]],
        ["MW", feb["mw"], buf["mw"]],
        ["dG_bind (WT)", feb["dG_bind_wt"] + " kcal/mol", buf["dG_bind_wt"] + " kcal/mol"],
        ["dG_bind (R702W)", feb["dG_bind_mut"] + " kcal/mol", buf["dG_bind_mut"] + " kcal/mol"],
        ["ddG (MUT-WT)", feb["ddG"] + " kcal/mol", buf["ddG"] + " kcal/mol"],
        ["Kd (WT)", feb["Kd_wt"], buf["Kd_wt"]],
        ["Kd (R702W)", feb["Kd_mut"], buf["Kd_mut"]],
        ["Mutation Effect", feb["mutation_effect"], buf["mutation_effect"]],
        ["Significance", feb["significance"], buf["significance"]],
    ]
    add_table(doc, headers, rows, header_color="27AE60")

    doc.add_heading("Key Interpretation", 2)
    interp = doc.add_paragraph()
    interp.add_run("- Bufadienolide predicted to bind much stronger than Febuxostat (requires experimental validation)\n").bold = True
    interp.add_run("- Bufadienolide shows mutation-resistant binding in silico\n").bold = True
    interp.add_run("- Febuxostat binding reduced ~50-fold in R702W (may reduce efficacy, requires experimental validation)\n")
    interp.add_run("- Bufadienolides are known NF-kB inhibitors - NOD2 signals through NF-kB\n")

    doc.add_page_break()

    # ========================================================================
    # SECTION 4: MM-GBSA vs FEP DISCREPANCY
    # ========================================================================
    doc.add_heading("4. MM-GBSA vs FEP Discrepancy", 1)

    critical = doc.add_paragraph()
    critical.add_run("CRITICAL FINDING: ").bold = True
    critical.add_run("MM-GBSA failed to accurately predict mutation effects. For Febuxostat, it predicted the opposite sign. For Bufadienolide, it overestimated the effect by 9x.")

    headers = ["Compound", "MM-GBSA ddG", "FEP ddG", "Direction"]
    rows = [[c["compound"], c["mmgbsa_ddg"] + " kcal/mol", c["fep_ddg"] + " kcal/mol", c["direction"]]
            for c in MMGBSA_FEP_COMPARISON]
    add_table(doc, headers, rows, header_color="E74C3C")

    doc.add_heading("Why FEP is Correct", 2)
    why = doc.add_paragraph()
    why.add_run("1. FEP is the gold standard for free energy calculations\n")
    why.add_run("2. MM-GBSA fails for charge-changing mutations (ARG+ -> TRP neutral)\n")
    why.add_run("3. MM-GBSA misses entropy contributions\n")
    why.add_run("4. 70-80 A allosteric effects require proper sampling\n")

    doc.add_heading("ISEF Narrative", 2)
    narrative = doc.add_paragraph()
    narrative.add_run(
        '"This demonstrates why rigorous free energy methods are essential for precision medicine, '
        'especially for charge-changing mutations at allosteric sites."'
    ).italic = True

    doc.add_page_break()

    # ========================================================================
    # SECTION 5: FILE PATHS
    # ========================================================================
    doc.add_heading("5. File Paths", 1)

    doc.add_paragraph("Base Directory: C:\\Users\\vasud\\nod2-screening-data\\")

    tree = doc.add_paragraph()
    tree.add_run("""
nod2-screening-data/
|-- receptor/
|   +-- human_NOD2_LRR.pdbqt           # NOD2 structure for docking
|
|-- data/
|   |-- natural_products_raw.csv       # Original screening library
|   +-- screening_results_final.csv    # All docking results
|
|-- PHASE_6/
|   |-- structures/
|   |   |-- febuxostat_docked.sdf
|   |   +-- natural_top_docked.sdf     # Contains CID_10120 (mislabeled)
|   +-- trajectories/                  # WT MD trajectories
|
|-- PHASE_A1_mutant_MD/
|   |-- trajectories/                  # Mutant MD (~58 GB)
|   +-- analysis/mmgbsa/               # MM-GBSA results
|
|-- fep_pmx/                           # Febuxostat FEP (80 total: 60 complex + 20 solvent)
|   |-- complex_WT/                    # 30 windows
|   |-- complex_MUT/                   # 30 windows
|   +-- solvent/                       # 20 windows
|
|-- fep_pmx_natural/                   # Bufadienolide FEP (60 total: 40 complex + 20 solvent)
|   |-- wt_complex/
|   |-- mut_complex/
|   +-- solvent/
|
+-- cross_validation/                  # DiffDock, Chai-1 validation
""")

    doc.add_page_break()

    # ========================================================================
    # SECTION 6: COMPOUND IDENTITY INVESTIGATION
    # ========================================================================
    doc.add_heading("6. Compound Identity Investigation", 1)

    warning = doc.add_paragraph()
    warning.add_run("CRITICAL CORRECTION: ").bold = True
    warning.add_run("The natural product was mislabeled throughout the project.")

    doc.add_heading("Correction Mapping", 2)
    old = COMPOUND_CORRECTION["old"]
    new = COMPOUND_CORRECTION["new"]

    headers = ["Property", "OLD (Wrong)", "NEW (Correct)"]
    rows = [
        ["PubChem CID", old["cid"], new["cid"]],
        ["Name", old["name"], new["common_name"]],
        ["IUPAC Name", "-", new["name"]],
        ["Type", old["type"], new["type"]],
        ["Formula", old["formula"], new["formula"]],
        ["MW", old["mw"], new["mw"]],
    ]
    add_table(doc, headers, rows, header_color="9B59B6")

    doc.add_heading("Root Cause", 2)
    doc.add_paragraph(COMPOUND_CORRECTION["root_cause"])

    doc.add_heading("Action Taken", 2)
    doc.add_paragraph(
        "All documentation updated to reflect correct compound identity (CID_10120, Bufadienolide). "
        "The science remains valid - only the compound label was incorrect."
    )

    doc.add_page_break()

    # ========================================================================
    # SECTION 7: STATISTICS SUMMARY
    # ========================================================================
    doc.add_heading("7. Statistics Summary", 1)

    headers = ["Metric", "Value"]
    rows = [
        ["Total compounds screened", STATISTICS["total_compounds_screened"]],
        ["Library breakdown", STATISTICS["library_breakdown"]],
        ["ML Model AUC-ROC", STATISTICS["ml_model_auc_roc"]],
        ["MD Simulations (WT)", STATISTICS["md_simulations_wt"]],
        ["MD Simulations (Mutant)", STATISTICS["md_simulations_mutant"]],
        ["Total simulation time", STATISTICS["total_simulation_time"]],
        ["FEP windows (Febuxostat)", STATISTICS["fep_windows_febuxostat"]],
        ["FEP windows (Bufadienolide)", STATISTICS["fep_windows_bufadienolide"]],
        ["Solvent windows", STATISTICS["solvent_windows"]],
        ["Total FEP windows", STATISTICS["total_fep_windows"]],
    ]
    add_table(doc, headers, rows)

    doc.add_page_break()

    # ========================================================================
    # SECTION 8: METHODS DETAIL
    # ========================================================================
    doc.add_heading("8. Methods Detail", 1)

    doc.add_heading("8.1 Virtual Screening", 2)
    vs = doc.add_paragraph()
    vs.add_run("- Software: GNINA v1.0 (GPU-accelerated)\n")
    vs.add_run("- Scoring: CNN affinity + CNN pose\n")
    vs.add_run("- Exhaustiveness: 32\n")
    vs.add_run("- Target: NOD2 LRR domain (AlphaFold-predicted structure, UniProt Q9HC29)\n")

    doc.add_heading("8.2 Machine Learning (NOD2-Scout)", 2)
    ml = doc.add_paragraph()
    ml.add_run("- Algorithm: Random Forest Classifier\n")
    ml.add_run("- Features: 2048-bit Morgan fingerprints + 15 physicochemical descriptors\n")
    ml.add_run("- Training: Top 10% (positive) vs bottom 10% (negative) by docking score\n")
    ml.add_run("- Validation: 5-fold cross-validation with scaffold split\n")
    ml.add_run("- Negative control: Label shuffling -> AUC = 0.50 (confirmed no spurious learning)\n")

    doc.add_heading("8.3 Molecular Dynamics", 2)
    md = doc.add_paragraph()
    md.add_run("- Software: OpenMM 8.0\n")
    md.add_run("- Force field: AMBER ff14SB (protein) + GAFF2 (ligand)\n")
    md.add_run("- Water: TIP3P, 10 A padding\n")
    md.add_run("- Temperature: 300 K (Langevin thermostat)\n")
    md.add_run("- Pressure: 1 atm (Monte Carlo barostat)\n")
    md.add_run("- Timestep: 2 fs with SHAKE constraints\n")
    md.add_run("- Duration: 20 ns per trajectory\n")

    doc.add_heading("8.4 Free Energy Perturbation", 2)
    fep = doc.add_paragraph()
    fep.add_run("- Software: OpenMM 8.2 with openmmtools alchemical factory\n")
    fep.add_run("- Ligand parameterization: OpenFF 2.1.0 (Sage) force field\n")
    fep.add_run("- Method: Alchemical FEP with soft-core potentials (alpha=0.5)\n")
    fep.add_run("- Windows: 20-30 per leg (electrostatics then sterics decoupling)\n")
    fep.add_run("- Sampling: 1 ns per window\n")
    fep.add_run("- Analysis: MBAR (pymbar 4.0)\n")
    fep.add_run("- Thermodynamic cycle: dG_bind = dG_solvent - dG_complex (negative = favorable binding)\n")

    doc.add_page_break()

    # ========================================================================
    # SECTION 9: NOVEL CONTRIBUTIONS
    # ========================================================================
    doc.add_heading("9. Novel Contributions", 1)

    novel_list = doc.add_paragraph()
    novel_list.add_run("1. First report of bufadienolide-NOD2 binding").bold = True
    novel_list.add_run(" - nobody has connected this compound class to NOD2\n\n")

    novel_list.add_run("2. FEP validation of NOD2 drug candidates").bold = True
    novel_list.add_run(" - first rigorous free energy study on NOD2 ligands\n\n")

    novel_list.add_run("3. Demonstration that MM-GBSA fails for allosteric charge-changing mutations").bold = True
    novel_list.add_run(" - important methodological finding\n\n")

    novel_list.add_run("4. Mutation-resistant drug discovery").bold = True
    novel_list.add_run(" - identified a compound unaffected by disease-causing mutation\n\n")

    novel_list.add_run("5. Personal relevance").bold = True
    novel_list.add_run(" - student has Crohn's disease with R702W mutation (N=1 precision medicine)\n")

    doc.add_page_break()

    # ========================================================================
    # SECTION 10: FUTURE DIRECTIONS
    # ========================================================================
    doc.add_heading("10. Future Directions", 1)

    future = doc.add_paragraph()
    future.add_run("1. Experimental validation").bold = True
    future.add_run(" - SPR or ITC binding assays for bufadienolide\n\n")

    future.add_run("2. Cell-based assays").bold = True
    future.add_run(" - NF-kB reporter assays in NOD2-expressing cells\n\n")

    future.add_run("3. Analog screening").bold = True
    future.add_run(" - test related bufadienolides for improved drug-likeness\n\n")

    future.add_run("4. Clinical translation").bold = True
    future.add_run(" - Febuxostat could enter clinical trials immediately (FDA-approved)\n")

    # ========================================================================
    # SAVE DOCUMENT
    # ========================================================================
    doc.save(OUTPUT_DOCX)
    print(f"Created: {OUTPUT_DOCX}")

    # Try to create PDF
    try:
        from docx2pdf import convert
        convert(str(OUTPUT_DOCX), str(OUTPUT_PDF))
        print(f"Created: {OUTPUT_PDF}")
    except ImportError:
        print("Note: docx2pdf not installed. Install with: pip install docx2pdf")
        print("      Or open the .docx in Word and export to PDF manually.")
    except Exception as e:
        print(f"PDF conversion failed: {e}")
        print("Open the .docx in Word and export to PDF manually.")

    return doc


if __name__ == "__main__":
    print("=" * 70)
    print("CREATING MASTER PROJECT SUMMARY")
    print("=" * 70)
    print(f"Output: {OUTPUT_DOCX}")
    print()

    doc = create_document()

    print()
    print("=" * 70)
    print("DONE")
    print("=" * 70)
